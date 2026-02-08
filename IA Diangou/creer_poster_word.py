#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script pour créer le poster IA Diangou au format Word
avec les couleurs et le formatage exact du template SIFSC 13
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERREUR: python-docx n'est pas installé.")
    print("Veuillez installer avec: pip install python-docx")
    exit(1)

# Créer le document
doc = Document()

# Définir les couleurs (RGB)
COULEUR_BLEU_FONCE = RGBColor(0, 51, 102)  # #003366
COULEUR_NOIR = RGBColor(0, 0, 0)
COULEUR_BLANC = RGBColor(255, 255, 255)

# Configuration de la page (format A0 pour poster)
section = doc.sections[0]
section.page_height = Inches(46.8)  # A0 height
section.page_width = Inches(33.1)   # A0 width
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)

def ajouter_paragraphe_centre(texte, taille_font, couleur_fond=None, couleur_texte=COULEUR_BLANC, gras=False):
    """Ajoute un paragraphe centré avec fond coloré"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(taille_font)
    run.font.bold = gras
    run.font.color.rgb = couleur_texte
    
    if couleur_fond:
        # Ajouter un fond coloré
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), couleur_fond.hex)
        pPr.append(shd)
    
    return p

def ajouter_titre_section(texte):
    """Ajoute un titre de section avec fond bleu"""
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COULEUR_NOIR
    return p

def ajouter_texte(texte, taille=16):
    """Ajoute du texte normal"""
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(taille)
    run.font.color.rgb = COULEUR_NOIR
    return p

# EN-TÊTE
ajouter_paragraphe_centre("SIFSC 13", 48, COULEUR_BLEU_FONCE, COULEUR_BLANC, True)
ajouter_paragraphe_centre("IA DIANGOU - INTELLIGENCE ARTIFICIELLE D'ENSEIGNEMENT IA DIANGOU - INTELLIGENCE ARTIFICIELLE D'ENSEIGNEMENT", 36, COULEUR_BLEU_FONCE, COULEUR_BLANC, True)
ajouter_paragraphe_centre("DIANGOU(a), Koolo Barry, DIANGOU(b), Koolo Barry", 24, COULEUR_BLEU_FONCE, COULEUR_BLANC, False)
ajouter_paragraphe_centre("(a) Institut de Recherche en Intelligence Artificielle. (b). Université de Conakry.", 20, COULEUR_BLEU_FONCE, COULEUR_BLANC, False)

doc.add_paragraph()  # Espace

# COLONNE GAUCHE - RESUMO
ajouter_titre_section("RESUMO")
ajouter_texte("IA Diangou est une plateforme d'intelligence artificielle éducative innovante conçue pour enseigner la langue française et répondre aux questions des apprenants de manière personnalisée et adaptative. Le système combine des technologies d'IA avancées (OpenAI, Hugging Face) avec une architecture web moderne (Flask, React, PostgreSQL) pour offrir une expérience d'apprentissage complète et accessible 24/7. La plateforme fonctionne en mode autonome (sans clé API) et en mode avancé (avec intégration d'API), permettant une utilisation flexible selon les besoins et ressources disponibles. L'innovation principale réside dans la capacité d'adaptation automatique du niveau d'enseignement du débutant au niveau avancé, avec une couverture complète de la grammaire, du vocabulaire, de la conjugaison et de la syntaxe française.")

doc.add_paragraph()  # Espace

# OBJETIVOS
ajouter_titre_section("OBJETIVOS")
ajouter_texte("Développer une intelligence artificielle capable d'enseigner le français de manière efficace, personnalisée et accessible. Les objectifs spécifiques incluent : (1) Enseignement adaptatif avec évaluation automatique du niveau de l'apprenant et progression du niveau débutant au niveau compétent, (2) Couverture complète de la langue française incluant grammaire, orthographe, vocabulaire, syntaxe, prononciation et conjugaison, (3) Accessibilité et flexibilité avec fonctionnement sans dépendance à des API payantes, (4) Gestion complète de la communauté éducative avec système d'inscription pour professeurs, parents et apprenants, et (5) Qualité pédagogique avec réponses simples, précises, bienveillantes et structurées.")

doc.add_paragraph()  # Espace

# METODOLOGIA
ajouter_titre_section("METODOLOGIA")
ajouter_texte("L'architecture technique utilise Flask (Python) pour le backend, React avec TypeScript pour le frontend, et PostgreSQL pour la base de données. Le système d'IA intègre trois modes de fonctionnement : un mode démo complet couvrant 52+ sujets pédagogiques sans dépendance externe, une intégration avec OpenAI API (GPT-3.5-turbo) pour des réponses avancées, et une alternative avec Hugging Face API. Le développement de l'IA repose sur un prompt système détaillé de 200+ lignes définissant le comportement pédagogique, un système de détection avec 52+ conditions pour identifier le type de question, et une fonction de conjugaison automatique supportant 8 verbes irréguliers principaux et 5 temps verbaux. Le processus de développement suit une approche structurée : conception et architecture, développement backend, développement frontend, tests et optimisation, puis déploiement et documentation.")

# Formule
p_formule = doc.add_paragraph()
run = p_formule.add_run("Architecture = Backend (Flask + PostgreSQL) + Frontend (React + TypeScript) + IA (Mode démo + OpenAI + Hugging Face)")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.color.rgb = COULEUR_NOIR
p_formule.add_run(" (1)")

doc.add_paragraph()  # Espace

# RESULTADOS
ajouter_titre_section("RESULTADOS")
ajouter_texte("Les résultats obtenus démontrent une plateforme fonctionnelle et complète. L'IA pédagogique fonctionne en mode démo couvrant 52+ sujets et en mode avancé avec intégration d'API. La conjugaison automatique détecte et conjugue les verbes en temps réel. La plateforme intègre une gestion complète des utilisateurs (professeurs, parents, apprenants), des formations et des cours. L'interface React moderne offre une expérience utilisateur optimale. Statistiquement, le projet compte environ 2992 lignes de code backend, 52+ sujets pédagogiques couverts, 7+ tables de base de données, 15+ endpoints API, avec un temps de réponse inférieur à 2 secondes.")

# Figure
p_fig = doc.add_paragraph()
run_fig = p_fig.add_run("Figura 1: Architecture et fonctionnalités de la plateforme IA Diangou")
run_fig.font.name = 'Times New Roman'
run_fig.font.size = Pt(16)
run_fig.font.bold = True
run_fig.font.color.rgb = COULEUR_NOIR

ajouter_texte("(a) Architecture technique du système")
ajouter_texte("(b) Interface utilisateur et expérience d'apprentissage")

ajouter_texte("Les tests fonctionnels confirment le bon fonctionnement de toutes les routes API, du système d'authentification, de la génération de réponses IA, et de la sauvegarde des conversations. Les tests d'intégration valident l'intégration backend-frontend, base de données, et API externes. Les tests utilisateurs confirment une interface intuitive, des réponses pédagogiques de qualité, et une expérience d'apprentissage positive.")

doc.add_paragraph()  # Espace

# TABLEAU (dans la colonne droite conceptuellement)
ajouter_titre_section("Tabela 1: Statistiques et métriques de performance de la plateforme IA Diangou")

# Créer le tableau
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'

# En-têtes
table.rows[0].cells[0].text = 'Métrique'
table.rows[0].cells[1].text = 'Valeur'
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)

# Données
donnees = [
    ('Lignes de code backend', '~2992'),
    ('Sujets pédagogiques couverts', '52+'),
    ('Tables de base de données', '7+'),
    ('Endpoints API', '15+'),
    ('Temps de réponse', '< 2 secondes'),
    ('Verbes irréguliers supportés', '8'),
    ('Temps verbaux supportés', '5')
]

for i, (metrique, valeur) in enumerate(donnees, 1):
    table.rows[i].cells[0].text = metrique
    table.rows[i].cells[1].text = valeur
    for cell in table.rows[i].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)

doc.add_paragraph()  # Espace

ajouter_texte("Les résultats attendus à court terme (3-6 mois) incluent l'extension de la conjugaison à 20+ verbes irréguliers, le développement d'un système d'exercices interactifs, l'atteinte de 100+ utilisateurs actifs, et la création de 50+ formations. À moyen terme (6-12 mois), les objectifs visent le support multilingue (anglais, espagnol, arabe), le développement d'applications mobiles (iOS, Android), l'atteinte de 1000+ utilisateurs actifs, et l'établissement de partenariats avec des écoles et universités.")

ajouter_texte("L'impact éducatif se manifeste par un accès 24/7 à l'éducation de qualité, un apprentissage personnalisé adapté au niveau, une réduction des coûts (gratuit ou à faible coût), et une flexibilité permettant l'apprentissage à son propre rythme. L'impact social contribue à la réduction des inégalités éducatives, à l'alphabétisation en français, à l'autonomisation des apprenants, et à l'accessibilité pour les populations rurales et défavorisées.")

doc.add_paragraph()  # Espace

# CONCLUSÕES
ajouter_titre_section("CONCLUSÕES")
ajouter_texte("IA Diangou démontre que l'intelligence artificielle peut être un outil puissant pour démocratiser l'accès à l'éducation de qualité. Le projet combine une IA pédagogique avancée adaptative et personnalisée, une plateforme éducative complète avec gestion de communauté, une architecture technique robuste utilisant des technologies modernes, et une accessibilité remarquable grâce au mode démo gratuit sans dépendance API. Les innovations principales incluent un système de détection de 52+ sujets pédagogiques, une conjugaison automatique intégrée, une adaptation automatique du niveau, et un mode démo complet sans dépendance externe. Les perspectives d'avenir visent l'extension des fonctionnalités pédagogiques, l'expansion géographique et linguistique, la contribution à la recherche en IA éducative, et un impact mesurable sur l'éducation et l'alphabétisation. L'avenir de l'éducation est entre nos mains, et l'IA est notre alliée pour créer un monde où chacun peut apprendre, grandir et réussir.")

doc.add_paragraph()  # Espace

# REFERÊNCIAS
ajouter_titre_section("REFERÊNCIAS")
ajouter_texte("[1] OPENAI. GPT-3.5 Turbo - Language Model. Platform OpenAI, https://platform.openai.com/docs/models/gpt-3-5, 2023.")
ajouter_texte("[2] HUGGING FACE. Transformers Library - State-of-the-art Machine Learning. Hugging Face, https://huggingface.co/docs/transformers, 2023.")
ajouter_texte("[3] FLASK DEVELOPMENT TEAM. Flask - Web Framework for Python. Pallets Projects, https://flask.palletsprojects.com/, 2023.")
ajouter_texte("[4] REACT TEAM. React - A JavaScript Library for Building User Interfaces. Meta, https://react.dev/, 2023.")
ajouter_texte("[5] POSTGRESQL GLOBAL DEVELOPMENT GROUP. PostgreSQL - Advanced Open Source Database. PostgreSQL, https://www.postgresql.org/, 2023.")

doc.add_paragraph()  # Espace

# APOIADORES
ajouter_paragraphe_centre("APOIADORES", 18, COULEUR_BLEU_FONCE, COULEUR_BLANC, True)
ajouter_paragraphe_centre("Projet IA Diangou - Intelligence Artificielle d'Enseignement", 14, None, COULEUR_BLEU_FONCE, False)
ajouter_paragraphe_centre("Institut de Recherche en Intelligence Artificielle", 14, None, COULEUR_BLEU_FONCE, False)
ajouter_paragraphe_centre("Université de Conakry", 14, None, COULEUR_BLEU_FONCE, False)
ajouter_paragraphe_centre("Département d'Informatique et Intelligence Artificielle", 14, None, COULEUR_BLEU_FONCE, False)

# Sauvegarder le document
nom_fichier = "Poster_IA_Diangou_SIFSC.docx"
doc.save(nom_fichier)
print(f"✅ Document Word créé avec succès : {nom_fichier}")
print("📄 Vous pouvez maintenant l'ouvrir dans Microsoft Word")

